#!/usr/bin/env python3
"""
将低空经济分析报告中的所有Markdown文件转换为Word格式
支持：标题层级、表格、加粗、列表、代码块、引用
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

REPORT_DIR = "/Users/kevinshi/Documents/workspace/codebase/shibit/xuanji/低空经济分析报告"
OUTPUT_DIR = "/Users/kevinshi/Documents/workspace/codebase/shibit/xuanji/低空经济分析报告/word输出"


def ensure_output_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_formatted_run(paragraph, text, bold=False, size=None, color=None, font_name=None):
    """添加格式化文本段落"""
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def parse_markdown_to_docx(md_content, doc):
    """解析Markdown内容并写入Word文档"""

    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_data = []
    table_alignments = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            if in_table and table_data:
                _write_table(doc, table_data, table_alignments)
                table_data = []
                table_alignments = []
                in_table = False
            i += 1
            continue

        # 表格行（以 | 开头）
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]

            # 分隔行（|---|）
            if re.match(r'^[\s\|:-]+$', stripped):
                table_alignments = []
                for c in stripped.split('|')[1:-1]:
                    if c.strip().startswith(':') and c.strip().endswith(':'):
                        table_alignments.append('center')
                    elif c.strip().endswith(':'):
                        table_alignments.append('right')
                    else:
                        table_alignments.append('left')
                i += 1
                continue

            # 数据行
            if not in_table:
                in_table = True
            table_data.append(cells)
            i += 1
            continue
        else:
            if in_table and table_data:
                _write_table(doc, table_data, table_alignments)
                table_data = []
                table_alignments = []
                in_table = False

        # 标题（# 开头）
        if stripped.startswith('#'):
            level = len(stripped.split(' ')[0])
            title_text = stripped.lstrip('#').strip()
            heading = doc.add_heading(title_text, level=min(level, 4))
            # 设置中文字体
            for run in heading.runs:
                run.font.name = 'SimHei'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            i += 1
            continue

        # 水平线（---）
        if re.match(r'^-{3,}$', stripped):
            # 添加一个带底部边框的段落作为分隔
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # 无序列表
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            # 处理加粗
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, text, font_name='SimSun')
            # 检查是否有子项
            j = i + 1
            while j < len(lines) and (lines[j].startswith('  - ') or lines[j].startswith('    - ') or lines[j].startswith('  * ') or lines[j].startswith('    * ')):
                sub_text = lines[j].strip()[2:].strip()
                sub_text = re.sub(r'\*\*(.+?)\*\*', r'\1', sub_text)
                sp = doc.add_paragraph(style='List Bullet 2')
                add_formatted_run(sp, sub_text, font_name='SimSun')
                j += 1
            if j > i + 1:
                i = j
                continue
            i += 1
            continue

        # 有序列表
        if re.match(r'^\d+[\.\、]', stripped):
            text = re.sub(r'^\d+[\.\、]\s*', '', stripped)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_run(p, text, font_name='SimSun')
            i += 1
            continue

        # 引用（> 开头）
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            # 添加引用竖线效果
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:left w:val="single" w:sz="24" w:space="8" w:color="2E75B6"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            run.font.name = 'SimSun'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            i += 1
            continue

        # 普通段落（处理加粗）
        text = stripped
        has_bold = '**' in text

        if has_bold:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    add_formatted_run(p, part[2:-2], bold=True, font_name='SimSun')
                else:
                    add_formatted_run(p, part, font_name='SimSun')
        else:
            p = doc.add_paragraph()
            add_formatted_run(p, text, font_name='SimSun')

        i += 1

    # 处理文档末尾的表格
    if in_table and table_data:
        _write_table(doc, table_data, table_alignments)


def _write_table(doc, table_data, alignments):
    """写入表格"""
    if not table_data:
        return

    rows = len(table_data)
    cols = max(len(row) for row in table_data)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格字体大小
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= cols:
                break
            cell = table.cell(row_idx, col_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = 'SimSun'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            # 对齐方式
            if alignments and col_idx < len(alignments):
                align_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                }
                p.alignment = align_map.get(alignments[col_idx], WD_ALIGN_PARAGRAPH.LEFT)

            # 表头加粗/底色（第一行）
            if row_idx == 0:
                run.bold = True
                set_cell_shading(cell, "2E75B6")
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()  # 表后空行


def convert_md_file(md_path, output_name=None):
    """转换单个Markdown文件为Word"""
    if output_name is None:
        output_name = os.path.splitext(os.path.basename(md_path))[0] + '.docx'

    output_path = os.path.join(OUTPUT_DIR, output_name)

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    parse_markdown_to_docx(md_content, doc)
    doc.save(output_path)
    return output_path


def main():
    ensure_output_dir()

    # 文件映射：md文件 -> docx文件名
    files_to_convert = [
        (os.path.join(REPORT_DIR, "关于低空经济发展态势及我们企业进军策略的分析报告.md"),
         "01-关于低空经济发展态势及我们企业进军策略的分析报告.docx"),
    ]

    # 添加政策文件
    policy_dir = os.path.join(REPORT_DIR, "政策文件")
    if os.path.exists(policy_dir):
        for fname in sorted(os.listdir(policy_dir)):
            if fname.endswith('.md'):
                md_path = os.path.join(policy_dir, fname)
                docx_name = fname.replace('.md', '.docx')
                files_to_convert.append((md_path, docx_name))

    print("=" * 60)
    print("  低空经济分析报告 → Word格式转换")
    print("=" * 60)
    print()

    for md_path, docx_name in files_to_convert:
        try:
            output_path = convert_md_file(md_path, docx_name)
            print(f"  ✅ {docx_name}")
        except Exception as e:
            print(f"  ❌ {docx_name}: {e}")

    print()
    print("-" * 60)
    print(f"  输出目录: {OUTPUT_DIR}")
    print(f"  共转换 {len(files_to_convert)} 个文件")
    print("=" * 60)


if __name__ == "__main__":
    main()
